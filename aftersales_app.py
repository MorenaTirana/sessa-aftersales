# AfterSalesApp - Portable v4.7.6 (tema Sessa + WIR/SPR + DOCX) — build pulita

import os
import sqlite3
import datetime as dt
from contextlib import contextmanager
import base64
from io import BytesIO
import pandas as pd


import streamlit as st
from PIL import Image
def require_auth():
    """Blocca l'app finché non inserisci user/pass. Legge APP_USER/APP_PASS da env o secrets."""
    USER = os.environ.get("APP_USER") or st.secrets.get("APP_USER", None)
    PASS = os.environ.get("APP_PASS") or st.secrets.get("APP_PASS", None)
    # Se non sono impostate, niente login (utile in locale)
    if not USER or not PASS:
        return
    ok = st.session_state.get("auth_ok", False)
    if not ok:
        with st.sidebar:
            st.markdown("### 🔐 Accesso")
            u = st.text_input("User", key="auth_user")
            p = st.text_input("Password", type="password", key="auth_pass")
            if st.button("Entra", key="auth_btn"):
                st.session_state.auth_ok = (u == USER and p == PASS)
                ok = st.session_state.auth_ok
        if not ok:
            st.warning("Inserisci credenziali per accedere.")
            st.stop()


# ====== pagina (va prima di qualunque chiamata Streamlit)
def _find_logo():
    candidates = [
        os.path.join(os.environ.get("AFTERSALES_DATA_ROOT", "./data"), "logo.png"),
        os.path.join(".", "data", "logo.png"),
        os.path.join(".", "logo.png"),
    ]
    for p in candidates:
        if os.path.exists(p):
            return p
    return None

LOGO = _find_logo()

def _logo_b64():
    try:
        if LOGO and os.path.exists(LOGO):
            with open(LOGO, "rb") as f:
                return base64.b64encode(f.read()).decode("utf-8")
    except Exception:
        pass
    return None

def _page_icon():
    try:
        if LOGO and os.path.exists(LOGO):
            return Image.open(LOGO)
    except Exception:
        pass
    return "⚓"

st.set_page_config(page_title="SESSA After Sales", page_icon=_page_icon(), layout="wide")


# ====== percorsi
DATA_ROOT = os.environ.get("AFTERSALES_DATA_ROOT", "./data")
DB_PATH   = os.environ.get("AFTERSALES_DB", os.path.join(DATA_ROOT, "db", "aftersales.db"))
BACKUPS   = os.path.join(DATA_ROOT, "backups")
UPLOADS   = os.path.join(DATA_ROOT, "uploads")
EXPORTS   = os.path.join(DATA_ROOT, "exports")
for d in (os.path.dirname(DB_PATH), BACKUPS, UPLOADS, EXPORTS):
    os.makedirs(d, exist_ok=True)


# ====== CSS tema Sessa
CUSTOM_CSS = """
<style>
:root{
  --sea:   #2f9ec5;   /* azzurro mare */
  --deep:  #0a1f44;   /* blu scuro testi */
  --rope:  #0b3b6e;   /* blu “corda” */
  --off:   #f7f9fb;   /* bianco sporco */
  --sessa: #3E79B3;   /* BLU SESSA */
  --navy:  #0b2a4a;   /* blu navy per bottoni */
}

/* Sfondo e font */
body, .stApp, .main .block-container{
  background: var(--off) !important;
  font-family: 'Times New Roman', Times, serif !important;
}

/* Titoli */
h1, h2, h3, .hero-title{
  color: var(--sessa) !important;
  font-weight: 700 !important;
  letter-spacing: .5px;
}

/* Sidebar blu Sessa */
section[data-testid="stSidebar"],
div[data-testid="stSidebar"],
*[data-testid="stSidebar"]{
  background: var(--sessa) !important;
}
section[data-testid="stSidebar"] > div,
div[data-testid="stSidebar"] > div{
  background: transparent !important;
  padding: 10px 10px 18px 10px;
  border-right: none !important;
}
*[data-testid="stSidebar"] *{ color:#fff !important; }
*[data-testid="stSidebar"] a{ color:#fff !important; }

/* Logo nella sidebar */
.sidebar-brand img{
  border-radius: 10px;
  margin-bottom: 10px;
  box-shadow: 0 4px 12px rgba(0,0,0,.25);
}

/* ===== MENU (sidebar) — LINK, no radio ===== */

/* Titolo MENU più in basso */
[data-testid="stSidebar"] .menu-title{
  font-size: 20px; font-weight: 800; letter-spacing:.3px;
  margin: 32px 0 10px 0 !important;
}

/* lista menu */
[data-testid="stSidebar"] .menu-list{
  display: flex; flex-direction: column; gap: 4px;
  margin-top: 6px;
}

/* voce menu (tutte light di default) */
[data-testid="stSidebar"] .menu-link{
  display: block;
  padding: 6px 4px;
  color: #fff !important;
  text-decoration: none !important;
  font-size: 14px !important;
  font-weight: 300 !important;      /* light */
  background: transparent !important;
  border-radius: 6px;
}

/* hover */
[data-testid="stSidebar"] .menu-link:hover{
  text-decoration: underline !important;
  text-underline-offset: 3px !important;
}

/* voce attiva: bold + underline */
[data-testid="stSidebar"] .menu-link.active{
  font-weight: 800 !important;      /* bold */
  text-decoration: underline !important;
  text-underline-offset: 3px !important;
  text-decoration-thickness: 2px !important;
}

/* nascondi eventuali radio residui */
[data-testid="stSidebar"] [role="radio"],
[data-testid="stSidebar"] [data-baseweb="radio"]{
  display: none !important;
}

/* Bottoni blu navy */
.stButton>button{
  background: var(--navy) !important;
  color:#fff !important;
  border: 0 !important;
  border-radius: 10px !important;
  font-weight: 700 !important;
  display: inline-flex !important;
  align-items: center !important;
  width: fit-content !important;
  white-space: nowrap !important;
  padding: 10px 16px !important;
}

/* --- Aggiunta: pulsanti del MENU nella sidebar resi come link --- */
[data-testid="stSidebar"] .stButton>button{
  background: transparent !important;
  color:#fff !important;
  border:0 !important;
  padding:6px 4px !important;
  text-align:left !important;
  width:100% !important;
  border-radius:6px !important;
  font-weight:300 !important;   /* light */
}

[data-testid="stSidebar"] .stButton>button:hover{
  text-decoration: underline !important;
  text-underline-offset: 3px !important;
}

/* Manteniamo bold + underline per la voce attiva (renderizzata come testo, non pulsante) */
[data-testid="stSidebar"] .menu-link.active{
  font-weight: 800 !important;
  text-decoration: underline !important;
  text-underline-offset: 3px !important;
  text-decoration-thickness: 2px !important;
}

/* Input box */
.stTextInput input, .stNumberInput input, .stDateInput input, .stTextArea textarea, [data-baseweb="select"] > div {
  background: #ffffff !important;
  border: 1px solid #d0d6df !important;
  border-radius: 8px !important;
  box-shadow: none !important;
}
.stTextInput input:focus, .stNumberInput input:focus, .stDateInput input:focus,
.stTextArea textarea:focus, [data-baseweb="select"] > div:focus-within {
  outline: none !important; border-color: #3E79B3 !important; box-shadow: 0 0 0 2px rgba(62,121,179,.15) !important;
}

/* Card bianche */
.card{ background:#ffffff !important; border-radius:12px !important; padding:14px 16px !important; box-shadow:0 2px 8px rgba(0,0,0,.05) !important; }

/* Card trasparente (per il submit / storico immediato) */
.card.no-bg{
  background: transparent !important;
  box-shadow: none !important;
  padding: 0 !important;
  margin: 0 !important;
}
.card.no-bg .stButton>button{ margin: 0 !important; }

/* Header blu con titolo bianco (grande accanto al logo) */
.brand-bar{
  background: var(--sessa); color:#fff;
  border-radius: 12px; padding: 10px 16px; margin: 0 0 14px 0;
  box-shadow: 0 6px 20px rgba(0,0,0,.12);
}
.brand-header{ display:flex; align-items:center; gap:12px; }
.brand-logo{
  width:72px; height:auto; border-radius:10px; border:2px solid rgba(255,255,255,.85);
  box-shadow:0 4px 12px rgba(0,0,0,.25);
}
.brand-text h1{ margin:0; padding:0; color:#fff !important; }
.brand-text small{ display:block; margin-top:2px; color:#fff !important; opacity:.95; }
@media (max-width:680px){ .brand-logo{ width:60px; } }

/* Elimina linee/divider ovunque (anche quelle sottilissime) */
hr, .stDivider, div[role="separator"], *[data-testid="stHorizontalRule"]{
  display:none !important; height:0 !important; border:0 !important; opacity:0 !important;
}
main .block-container div[aria-hidden="true"]{ display:none !important; }
main .block-container div[style*="border-top: 1px"],
main .block-container div[style*="border-bottom: 1px"]{ border-top:0 !important; border-bottom:0 !important; }
main .block-container div[style*="height: 1px"]{ height:0 !important; }
div[data-testid="stMarkdownContainer"]{ border:0 !important; box-shadow:none !important; outline:0 !important; }
</style>
"""
st.markdown(CUSTOM_CSS, unsafe_allow_html=True)

# ====== helpers DB
def get_conn():
    os.makedirs(os.path.dirname(DB_PATH), exist_ok=True)
    return sqlite3.connect(DB_PATH, check_same_thread=False)

def ensure_schema(conn: sqlite3.Connection):
    cur = conn.cursor()
    # WIR
    cur.execute("""CREATE TABLE IF NOT EXISTS wir(
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        created_at TEXT, dealer TEXT, boat TEXT,
        title TEXT, description TEXT, brand TEXT, serial TEXT,
        full_name TEXT, email TEXT, phone TEXT,
        boat_model TEXT, hull_serial TEXT,
        warranty_start TEXT, boat_location TEXT, onboard_contact TEXT
    )""")
    cur.execute("""CREATE TABLE IF NOT EXISTS wir_items(
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        wir_id INTEGER, req_index INTEGER,
        description TEXT, brand TEXT, item_serial TEXT, photo_paths TEXT
    )""")
    # SPR
    cur.execute("""CREATE TABLE IF NOT EXISTS spr(
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        created_at TEXT, dealer TEXT, boat TEXT,
        description TEXT, item_brand TEXT, item_serial TEXT,
        full_name TEXT, email TEXT, phone TEXT,
        boat_model TEXT, hull_serial TEXT,
        boat_location TEXT, onboard_contact TEXT
    )""")
    cur.execute("""CREATE TABLE IF NOT EXISTS spr_items(
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        spr_id INTEGER, req_index INTEGER,
        description TEXT, brand TEXT, item_serial TEXT, photo_paths TEXT
    )""")
    conn.commit()

def run_sql(conn, q, params=None):
    c = conn.cursor()
    c.execute(q, params or [])
    conn.commit()
    return c

def df_read(conn, table):
    try:
        import pandas as pd
        return pd.read_sql_query(f"SELECT * FROM {table} ORDER BY id DESC", conn)
    except Exception:
        import pandas as pd
        return pd.DataFrame()

# ====== DOCX (semplici)
def _ensure_folder(p): os.makedirs(p, exist_ok=True); return p
def _dmy(d): return d.strftime("%d_%m_%Y")

def export_wir_blank_docx(when_dt: dt.date, logo_path):
    from docx import Document
    from docx.shared import Inches
    folder = _ensure_folder(os.path.join(EXPORTS, "WIR"))
    name   = f"WIR - {_dmy(when_dt)}.docx"; path = os.path.join(folder, name)
    doc = Document()
    if logo_path: doc.add_picture(logo_path, width=Inches(1.2))
    doc.add_heading("WARRANTY INTERVENTION REQUEST (WIR)", level=1)
    doc.add_paragraph(when_dt.strftime("%d/%m/%Y"))
    doc.save(path)
    return path

def export_spr_blank_docx(when_dt: dt.date, logo_path):
    from docx import Document
    from docx.shared import Inches
    folder = _ensure_folder(os.path.join(EXPORTS, "SPR"))
    name   = f"SPR - {_dmy(when_dt)}.docx"; path = os.path.join(folder, name)
    doc = Document()
    if logo_path: doc.add_picture(logo_path, width=Inches(1.2))
    doc.add_heading("SPARE PARTS REQUEST (SPR)", level=1)
    doc.add_paragraph(when_dt.strftime("%d/%m/%Y"))
    doc.save(path)
    return path

def export_wir_docx(wir_id, when_dt, boat_model, header_map, items_for_doc, logo_path):
    from docx import Document
    from docx.shared import Inches
    folder = _ensure_folder(os.path.join(EXPORTS, "WIR"))
    name   = f"WIR {boat_model} - {_dmy(when_dt)}.docx"; path = os.path.join(folder, name)
    doc = Document()
    if logo_path: doc.add_picture(logo_path, width=Inches(1.2))
    doc.add_heading("WARRANTY INTERVENTION REQUEST (WIR)", level=1)
    for k,v in header_map.items(): doc.add_paragraph(f"{k}: {v}")
    doc.add_paragraph("------ RICHIESTE ------")
    for i, it in enumerate(items_for_doc, start=1):
        doc.add_paragraph(f"Richiesta {i}: {it['description']}")
        if it.get("brand"): doc.add_paragraph(f" Marchio: {it['brand']}")
        if it.get("item_serial"): doc.add_paragraph(f" Articolo/N.Serie: {it['item_serial']}")
    doc.save(path)
    return path

def export_spr_docx(spr_id, when_dt, boat_model, header_map, items_for_doc, logo_path):
    from docx import Document
    from docx.shared import Inches
    folder = _ensure_folder(os.path.join(EXPORTS, "SPR"))
    name   = f"SPR {boat_model} - {_dmy(when_dt)}.docx"; path = os.path.join(folder, name)
    doc = Document()
    if logo_path: doc.add_picture(logo_path, width=Inches(1.2))
    doc.add_heading("SPARE PARTS REQUEST (SPR)", level=1)
    for k,v in header_map.items(): doc.add_paragraph(f"{k}: {v}")
    doc.add_paragraph("------ RICHIESTE ------")
    for i, it in enumerate(items_for_doc, start=1):
        doc.add_paragraph(f"Richiesta {i}: {it['description']}")
        if it.get("brand"): doc.add_paragraph(f" Marchio: {it['brand']}")
        if it.get("item_serial"): doc.add_paragraph(f" Articolo/N.Serie: {it['item_serial']}")
    doc.save(path)
    return path

def export_after_sales_excel(conn, add_all_sheet: bool = True) -> bytes:
    """
    Esporta i dati correnti del DB in un Excel multi-foglio:
    - WIR, WIR_items
    - SPR, SPR_items
    - Materiali_in_garanzia (template vuoto)
    - Trasferte (template vuoto)
    - Barche_vendute (template vuoto)
    Opzionale: ALL_DATA (consolidato di tutti i fogli DB reali).
    """
    # leggi tabelle esistenti
    df_wir        = df_read(conn, "wir")
    df_wir_items  = df_read(conn, "wir_items")
    df_spr        = df_read(conn, "spr")
    df_spr_items  = df_read(conn, "spr_items")

    # assicura DataFrame anche se vuoti
    if df_wir is None:        df_wir = pd.DataFrame()
    if df_wir_items is None:  df_wir_items = pd.DataFrame()
    if df_spr is None:        df_spr = pd.DataFrame()
    if df_spr_items is None:  df_spr_items = pd.DataFrame()

    # template (colonne indicative, puoi rinominarle in futuro)
    cols_mat_gar = ["data", "codice", "descrizione", "quantita", "costo", "note"]
    cols_trip    = ["data", "tecnico", "barca", "luogo", "ore", "note"]
    cols_boats   = ["anno", "modello", "hull_serial", "cliente", "dealer", "note"]

    df_mat_gar   = pd.DataFrame(columns=cols_mat_gar)
    df_trip      = pd.DataFrame(columns=cols_trip)
    df_boats     = pd.DataFrame(columns=cols_boats)

    output = BytesIO()
    with pd.ExcelWriter(output, engine="xlsxwriter") as writer:
        df_wir.to_excel(writer, index=False, sheet_name="WIR")
        df_wir_items.to_excel(writer, index=False, sheet_name="WIR_items")
        df_spr.to_excel(writer, index=False, sheet_name="SPR")
        df_spr_items.to_excel(writer, index=False, sheet_name="SPR_items")

        df_mat_gar.to_excel(writer, index=False, sheet_name="Materiali_in_garanzia")
        df_trip.to_excel(writer, index=False, sheet_name="Trasferte")
        df_boats.to_excel(writer, index=False, sheet_name="Barche_vendute")

        if add_all_sheet:
            # ALL_DATA = unione colonne (schema largo) + colonna 'area'
            frames = []
            if not df_wir.empty:
                t = df_wir.copy(); t["area"] = "WIR"; frames.append(t)
            if not df_wir_items.empty:
                t = df_wir_items.copy(); t["area"] = "WIR_items"; frames.append(t)
            if not df_spr.empty:
                t = df_spr.copy(); t["area"] = "SPR"; frames.append(t)
            if not df_spr_items.empty:
                t = df_spr_items.copy(); t["area"] = "SPR_items"; frames.append(t)

            if frames:
                # allinea le colonne per union, riempiendo quelle mancanti
                all_cols = sorted(set().union(*[set(f.columns) for f in frames] | {"area"}))
                frames = [f.reindex(columns=all_cols) for f in frames]
                df_all = pd.concat(frames, ignore_index=True)
            else:
                df_all = pd.DataFrame(columns=["area"])

            df_all.to_excel(writer, index=False, sheet_name="ALL_DATA")

    output.seek(0)
    return output.getvalue()

# ====== piccoli helper UI
@contextmanager
def card(cls: str = ""):
    st.markdown(f'<div class="card {cls}">', unsafe_allow_html=True)
    try:
        yield
    finally:
        st.markdown('</div>', unsafe_allow_html=True)


# ====== SEZIONI
def header():
    img64 = _logo_b64()
    logo_html = f'<img src="data:image/png;base64,{img64}" class="brand-logo"/>' if img64 else ""
    st.markdown(
        f"""
        <div class="brand-bar">
          <div class="brand-header">
            {logo_html}
            <div class="brand-text">
              <h1>SESSA AFTER SALES</h1>
              <small>After Sales Dashboard</small>
            </div>
          </div>
        </div>
        """,
        unsafe_allow_html=True,
    )

def section_wir(conn):
    st.subheader("⚓ Warranty Intervention Requests (WIR)")

    if "wir_rows" not in st.session_state:
        st.session_state.wir_rows = [0]

    # Modulo WIR vuoto
    with st.expander("Scarica modulo WIR vuoto (DOCX)", expanded=False):
        d = st.date_input("Data modulo", dt.date.today(), key="wir_blank_date")
        if st.button("Scarica WIR", key="wir_blank_btn"):
            p = export_wir_blank_docx(d, LOGO)
            with open(p, "rb") as f:
                st.download_button("Download", f, file_name=os.path.basename(p), key="wir_blank_dl")

    # --- DATI
    st.markdown("### DATI")
    with card():
        c = st.columns(4)
        data = c[0].date_input("Data", dt.date.today(), key="wir_date")
        full_name = c[1].text_input("Nome & Cognome *", key="wir_fullname")
        dealer = c[2].text_input("Dealer *", key="wir_dealer")
        email = c[3].text_input("E-mail *", key="wir_email")

        c2 = st.columns(4)
        phone = c2[0].text_input("Cellulare *", key="wir_phone")
        boat_model = c2[1].text_input("Modello di barca *", key="wir_boat_model")
        hull_serial = c2[2].text_input("Matricola nr *", key="wir_hull")
        warranty_start = c2[3].date_input("Data attivazione garanzia *", dt.date.today(), key="wir_wstart")

        c3 = st.columns(2)
        boat_location = c3[0].text_input("Locazione barca", key="wir_loc")
        onboard_contact = c3[1].text_input("Contatto a bordo", key="wir_onboard")

    # --- RICHIESTE
    st.markdown("### RICHIESTE")
    with card():
        blocks = []
        for i, _ in enumerate(st.session_state.wir_rows, start=1):
            st.markdown("<div class='req-block'>", unsafe_allow_html=True)
            with st.expander(f"Richiesta {i}", expanded=True):
                desc = st.text_area("Descrizione *", key=f"wir_desc_{i}")
                up = st.file_uploader("Foto (una o più)", type=["png","jpg","jpeg"],
                                      accept_multiple_files=True, key=f"wir_ph_{i}")
                cc2 = st.columns(2)
                brand = cc2[0].text_input("Marchio", key=f"wir_brand_{i}")
                item_serial = cc2[1].text_input("Articolo / N. Serie", key=f"wir_item_{i}")
                blocks.append(dict(description=desc, uploads=up, brand=brand, item_serial=item_serial))
            st.markdown("</div>", unsafe_allow_html=True)

        cta = st.columns([1, 1, 6])
        with cta[0]:
            if st.button("➕ Aggiungi richiesta", key="wir_add_bottom"):
                st.session_state.wir_rows.append(len(st.session_state.wir_rows))
        with cta[1]:
            if len(st.session_state.wir_rows) > 1 and st.button("➖ Rimuovi ultima", key="wir_del_bottom"):
                st.session_state.wir_rows.pop()

        # SUBMIT (centrato, senza card bianco)
        with card("no-bg"):
            _, c2, _ = st.columns([1, 1, 1])
            with c2:
                submitted = st.button("✅ Salva & Genera Richiesta", key="wir_submit")
            if submitted:
                if not (full_name and dealer and email and phone and boat_model and hull_serial):
                    st.error("Compila i campi contrassegnati con *.")
                    st.stop()

                cur = run_sql(conn, """INSERT INTO wir(created_at,dealer,boat,title,description,brand,serial,
                                 full_name,email,phone,boat_model,hull_serial,warranty_start,boat_location,onboard_contact)
                                 VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
                              [data.isoformat(), dealer, boat_model, "", "", "", "",
                               full_name, email, phone, boat_model, hull_serial, warranty_start.isoformat(),
                               boat_location, onboard_contact])
                wir_id = cur.lastrowid

                base_dir = os.path.join(UPLOADS, "wir", str(wir_id)); os.makedirs(base_dir, exist_ok=True)
                items_for_doc = []
                for idx, b in enumerate(blocks, start=1):
                    saved = []
                    for f in (b["uploads"] or []):
                        fn = f"req{idx}_{f.name}"
                        p = os.path.join(base_dir, fn)
                        with open(p, "wb") as out: out.write(f.read())
                        saved.append(p)
                    run_sql(conn, """INSERT INTO wir_items(wir_id,req_index,description,brand,item_serial,photo_paths)
                                     VALUES(?,?,?,?,?,?)""",
                           [wir_id, idx, b["description"], b["brand"], b["item_serial"], ";".join(saved)])
                    items_for_doc.append({"description": b["description"], "brand": b["brand"],
                                          "item_serial": b["item_serial"], "photos": saved})

                header_map = {
                    "Nome & Cognome": full_name, "Dealer": dealer, "E-mail": email, "Cellulare": phone,
                    "Modello barca": boat_model, "Matricola": hull_serial,
                    "Data attivazione garanzia": warranty_start.strftime("%d/%m/%Y"),
                    "Locazione barca": boat_location, "Contatto a bordo": onboard_contact
                }
                docx_path = export_wir_docx(wir_id, data, boat_model, header_map, items_for_doc, LOGO)
                st.success("Grazie per averci contattati. Il team Sessa ha ricevuto la sua pratica e la contatterà per assisterla ulteriormente. Cordiali saluti! Sessa International Team")
                with open(docx_path, "rb") as f:
                    st.download_button("📄 Scarica modulo WIR", f, file_name=os.path.basename(docx_path), key="wir_doc_dl")
                st.session_state.wir_rows = [0]

    # LINK STORICO (senza card bianco immediato)
    with card("no-bg"):
        st.markdown("### WIR")
        dfw = df_read(conn, "wir")
        if dfw.empty:
            st.info("Nessuna WIR presente.")
        else:
            with card():
                st.dataframe(dfw, use_container_width=True)

def section_spr(conn):
    st.subheader("🛠️ Spare Parts Request (SPR)")

    if "spr_rows" not in st.session_state:
        st.session_state.spr_rows = [0]

    # Modulo SPR vuoto
    with st.expander("Scarica modulo SPR vuoto (DOCX)", expanded=False):
        d = st.date_input("Data modulo", dt.date.today(), key="spr_blank_date")
        if st.button("Scarica SPR vuota", key="spr_blank_btn"):
            p = export_spr_blank_docx(d, LOGO)
            with open(p, "rb") as f:
                st.download_button("Download", f, file_name=os.path.basename(p), key="spr_blank_dl")

    # --- DATI
    st.markdown("### DATI")
    with card():
        c = st.columns(4)
        data = c[0].date_input("Data", dt.date.today(), key="spr_date")
        full_name = c[1].text_input("Nome & Cognome *", key="spr_fullname")
        dealer = c[2].text_input("Dealer *", key="spr_dealer")
        email = c[3].text_input("E-mail *", key="spr_email")

        c2 = st.columns(4)
        phone = c2[0].text_input("Cellulare *", key="spr_phone")
        boat_model = c2[1].text_input("Modello di barca *", key="spr_boat_model")
        hull_serial = c2[2].text_input("Matricola nr *", key="spr_hull")
        _gap = c2[3].empty()

        c3 = st.columns(2)
        boat_location = c3[0].text_input("Locazione barca", key="spr_loc")
        onboard_contact = c3[1].text_input("Contatto a bordo", key="spr_onboard")

    # --- RICHIESTE
    st.markdown("### RICHIESTE")
    with card():
        blocks = []
        for i, _ in enumerate(st.session_state.spr_rows, start=1):
            st.markdown("<div class='req-block'>", unsafe_allow_html=True)
            with st.expander(f"Richiesta {i}", expanded=True):
                desc = st.text_area("Descrizione *", key=f"spr_desc_{i}")
                up = st.file_uploader("Foto (una o più)", type=["png","jpg","jpeg"],
                                      accept_multiple_files=True, key=f"spr_ph_{i}")
                cc2 = st.columns(2)
                brand = cc2[0].text_input("Marchio", key=f"spr_brand_{i}")
                item_serial = cc2[1].text_input("Articolo / N. Serie", key=f"spr_item_{i}")
                blocks.append(dict(description=desc, uploads=up, brand=brand, item_serial=item_serial))
            st.markdown("</div>", unsafe_allow_html=True)

        cta = st.columns([1, 1, 6])
        with cta[0]:
            if st.button("➕ Aggiungi richiesta", key="spr_add_bottom"):
                st.session_state.spr_rows.append(len(st.session_state.spr_rows))
        with cta[1]:
            if len(st.session_state.spr_rows) > 1 and st.button("➖ Rimuovi ultima", key="spr_del_bottom"):
                st.session_state.spr_rows.pop()

        # SUBMIT (centrato, senza card bianco)
        with card("no-bg"):
            _, c2, _ = st.columns([1, 1, 1])
            with c2:
                submitted = st.button("✅ Salva & Genera Richiesta", key="spr_submit")
            if submitted:
                if not (full_name and dealer and email and phone and boat_model and hull_serial):
                    st.error("Compila i campi contrassegnati con *.")
                    st.stop()

                cur = run_sql(conn, """INSERT INTO spr(created_at,dealer,boat,description,item_brand,item_serial,
                                 full_name,email,phone,boat_model,hull_serial,boat_location,onboard_contact)
                                 VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?)""",
                              [data.isoformat(), dealer, boat_model, "", "", "",
                               full_name, email, phone, boat_model, hull_serial, boat_location, onboard_contact])
                spr_id = cur.lastrowid

                base_dir = os.path.join(UPLOADS, "spr", str(spr_id)); os.makedirs(base_dir, exist_ok=True)
                items_for_doc = []
                for idx, b in enumerate(blocks, start=1):
                    saved = []
                    for f in (b["uploads"] or []):
                        fn = f"req{idx}_{f.name}"
                        p = os.path.join(base_dir, fn)
                        with open(p, "wb") as out: out.write(f.read())
                        saved.append(p)
                    run_sql(conn, """INSERT INTO spr_items(spr_id,req_index,description,brand,item_serial,photo_paths)
                                     VALUES(?,?,?,?,?,?)""",
                           [spr_id, idx, b["description"], b["brand"], b["item_serial"], ";".join(saved)])
                    items_for_doc.append({"description": b["description"], "brand": b["brand"],
                                          "item_serial": b["item_serial"], "photos": saved})

                header_map = {
                    "Nome & Cognome": full_name, "Dealer": dealer, "E-mail": email, "Cellulare": phone,
                    "Modello barca": boat_model, "Matricola": hull_serial,
                    "Locazione barca": boat_location, "Contatto a bordo": onboard_contact
                }
                docx_path = export_spr_docx(spr_id, data, boat_model, header_map, items_for_doc, LOGO)
                st.success("Grazie per averci contattati. Il team Sessa ha ricevuto la sua pratica e la contatterà per assisterla ulteriormente. Cordiali saluti! Sessa International Team")
                with open(docx_path, "rb") as f:
                    st.download_button("📄 Scarica modulo SPR", f, file_name=os.path.basename(docx_path), key="spr_doc_dl")
                st.session_state.spr_rows = [0]

    # LINK STORICO (senza card bianco immediato)
    with card("no-bg"):
        st.markdown("### SPR")
        dfs = df_read(conn, "spr")
        if dfs.empty:
            st.info("Nessuna SPR presente.")
        else:
            with card():
                st.dataframe(dfs, use_container_width=True)

# (sezioni segnaposto – puoi lasciarle così)
def section_clienti(conn):
    st.subheader("CLIENTI")
    with card(): st.write("Anagrafica Clienti.")

def section_dealer(conn):
    st.subheader("DEALER")
    with card(): st.write("Anagrafica Dealer.")

def section_ddt033(conn):
    st.subheader("DOCUMENTI GARANZIE")
    with card(): st.write("Documenti Garanzie.")

def section_ddt006(conn):
    st.subheader("DOCUMENTI VENDITE")
    with card(): st.write("Documenti Vendite.")

def section_quotes_invoices(conn):
    st.subheader("€ RIPARAZIONI")
    with card(): st.write("Preventivi / Fatture.")

def section_trips(conn):
    st.subheader("TRASFERTE")
    with card(): st.write("Trasferte.")

def section_boats(conn):
    st.subheader("BARCHE")
    with card(): st.write("Archivio barche.")

def section_reports(conn):
    st.subheader("REPORT")

    with card():
        st.markdown("### 📤 Esporta Excel dai dati correnti")
        c1, c2 = st.columns(2)
        with c1:
            add_all = st.checkbox("Aggiungi foglio consolidato 'ALL_DATA'", value=True)
        with c2:
            file_name = f"AFTER_SALES_REPORT_{dt.date.today().strftime('%Y%m%d')}.xlsx"

        if st.button("Genera file Excel"):
            xbytes = export_after_sales_excel(conn, add_all_sheet=add_all)
            st.download_button("⬇️ Scarica Excel", data=xbytes, file_name=file_name, mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", use_container_width=True)

    with card():
        st.markdown("### 📥 Importa / Visualizza Excel")
        up = st.file_uploader("Carica un file .xlsx (il tuo foglio After Sales)", type=["xlsx"])
        if up is not None:
            try:
                xls = pd.ExcelFile(up)
                sheets = xls.sheet_names
                st.success(f"File caricato. Fogli trovati: {', '.join(sheets)}")

                view = st.radio("Come vuoi visualizzare?", ["Per foglio", "Tutti i fogli"], horizontal=True)
                if view == "Per foglio":
                    name = st.selectbox("Scegli il foglio", sheets)
                    df = pd.read_excel(xls, sheet_name=name)
                    st.dataframe(df, use_container_width=True)
                else:
                    for name in sheets:
                        st.markdown(f"**{name}**")
                        df = pd.read_excel(xls, sheet_name=name)
                        st.dataframe(df, use_container_width=True)
                        st.markdown("---")

                # Ridai la possibilità di riscaricare il file caricato (comodo per conferma)
                st.download_button("⬇️ Scarica il file caricato", data=up.getvalue(),
                                   file_name=getattr(up, "name", "dati_aftersales.xlsx"),
                                   mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
            except Exception as e:
                st.error(f"Impossibile leggere l'Excel: {e}")
        else:
            st.info("Carica un Excel per visualizzarlo qui sopra.")

def section_settings(conn):
    st.subheader("IMPOSTAZIONI")
    st.code(f"DATA_ROOT = {DATA_ROOT}\nDB_PATH   = {DB_PATH}")
    with card():
        if st.button("💾 Backup DB adesso"):
            ts = dt.datetime.now().strftime("%Y%m%d-%H%M%S")
            dst = os.path.join(BACKUPS, f"aftersales_{ts}.db")
            with open(DB_PATH, "rb") as src, open(dst, "wb") as out:
                out.write(src.read())
            st.success("Backup creato.")

def section_storico(conn):
    st.subheader("STORICO PRATICHE")
    with card():
        st.write("Storico pratiche (WIR + SPR)")
        import pandas as pd
        dfw = df_read(conn, "wir"); dfs = df_read(conn, "spr")
        if dfw is not None and not dfw.empty: st.markdown("### WIR"); st.dataframe(dfw, use_container_width=True)
        if dfs is not None and not dfs.empty: st.markdown("### SPR"); st.dataframe(dfs, use_container_width=True)


# ====== routing
PAGES = {
    "WIR": section_wir,
    "SPR": section_spr,
    "CLIENTI": section_clienti,
    "DEALER": section_dealer,
    "DOCUMENTI GARANZIE": section_ddt033,
    "DOCUMENTI VENDITE": section_ddt006,
    "€ RIPARAZIONI": section_quotes_invoices,
    "TRASFERTE": section_trips,
    "BARCHE": section_boats,
    "REPORT": section_reports,
    "IMPOSTAZIONI": section_settings,
    "STORICO PRATICHE": section_storico,
}

def main():
    # sidebar: logo
    if LOGO and os.path.exists(LOGO):
        st.sidebar.image(LOGO, use_container_width=True)

    require_auth()


    # ------ MENU con pulsanti (niente nuove schede) ------
    page_keys = list(PAGES.keys())
    default_page = page_keys[0]

    # Leggi/sincronizza pagina corrente (nuove/vecchie API)
    try:
        qp_now = st.query_params
        current = qp_now.get("nav", st.session_state.get("nav", default_page))
    except Exception:
        qp_now = st.experimental_get_query_params()
        current = qp_now.get("nav", [st.session_state.get("nav", default_page)])[0]

    if isinstance(current, list):
        current = current[0]
    if current not in PAGES:
        current = default_page
    st.session_state.nav = current

    # Titolo + voci
    st.sidebar.markdown('<div class="menu-title">MENU</div>', unsafe_allow_html=True)
    menu_box = st.sidebar.container()
    for page in page_keys:
        if page == current:
            # voce attiva come testo (bold+underline via CSS .menu-link.active)
            menu_box.markdown(f'<div class="menu-link active">{page}</div>', unsafe_allow_html=True)
        else:
            # altre voci come pulsanti che aggiornano i query params e rerun
            if menu_box.button(page, key=f"navbtn_{page}"):
                st.session_state.nav = page
                try:
                    st.query_params.update({"nav": page})          # nuove API
                except Exception:
                    st.experimental_set_query_params(nav=page)     # fallback
                st.rerun()
    # ------------------------------------

    # intestazione pagina
    header()

    # DB + routing
    conn = get_conn()
    try:
        ensure_schema(conn)
        PAGES[st.session_state.nav](conn)
    finally:
        conn.close()


if __name__ == "__main__":
    main()
